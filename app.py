# -*- coding: utf-8 -*-
import os
import fitz
import streamlit as st
import concurrent.futures
from docx import Document
from openai import OpenAI

# ========================
# 配置区域（请替换API密钥）⚠️
# ========================
DEEPSEEK_API_KEY = "sk-cba0efd438ce4ba8a548bbff7a9c80f1"
API_BASE_URL = "https://api.deepseek.com"

# ========================
# 初始化服务
# ========================
client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=API_BASE_URL)

# ========================
# 自定义样式
# ========================
def inject_custom_css():
    st.markdown("""
    <style>
    :root {
        --primary: #003366;
        --secondary: #E31837;
    }

    .header {
        background: linear-gradient(135deg, var(--primary), #004080);
        padding: 2rem;
        color: white;
        border-radius: 15px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }

    .security-alert {
        background: #fff3cd;
        border-left: 5px solid #ffc107;
        padding: 1rem;
        margin: 1rem 0;
        border-radius: 4px;
    }

    .stDownloadButton>button {
        background-color: var(--primary) !important;
        color: white !important;
        border-radius: 20px;
    }

    .stProgress > div > div > div {
        background-color: var(--primary) !important;
    }
    </style>
    """, unsafe_allow_html=True)

# ========================
# 核心功能函数
# ========================
def extract_pdf_text(file):
    """提取PDF文本内容"""
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return "\n".join(page.get_text("text") for page in doc)
    except Exception as e:
        st.error(f"PDF解析失败: {str(e)}")
        return ""

def generate_word_file(text, filename, action):
    """生成Word文档"""
    try:
        doc = Document()
        doc.add_heading(f"{filename} - {action}", 0)
        doc.add_paragraph(text)
        output_path = f"{filename}_{action}.docx"
        doc.save(output_path)
        return output_path
    except Exception as e:
        st.error(f"生成Word文件失败: {str(e)}")
        return None

def get_ai_response(prompt):
    """获取AI回复"""
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{
                "role": "system",
                "content": "你是东吴证券专业金融分析助手，提供严谨专业的分析报告"
            }, {
                "role": "user", 
                "content": prompt
            }],
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"AI服务请求失败: {str(e)}")
        return ""

# ========================
# 界面模块
# ========================
def show_chat_interface():
    """聊天对话界面"""
    st.subheader("💬 智能金融问答")
    
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    
    # 显示聊天记录
    chat_container = st.container(height=500)
    with chat_container:
        for msg in st.session_state.chat_history:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
    
    # 输入处理
    if prompt := st.chat_input("输入您的金融问题..."):
        st.session_state.chat_history.append({"role": "user", "content": prompt})
        
        with st.spinner("正在生成专业回复..."):
            response = get_ai_response(prompt)
            st.session_state.chat_history.append({"role": "assistant", "content": response})
        
        st.rerun()

def show_document_interface():
    """文档处理界面"""
    st.subheader("📊 智能文档分析")
    
    # 文件上传
    uploaded_files = st.file_uploader(
        "上传PDF文档（支持多文件）",
        type=["pdf"],
        accept_multiple_files=True,
        help="单个文件最大50MB"
    )
    
    if uploaded_files:
        # 处理选项
        with st.expander("⚙️ 处理设置", expanded=True):
            col1, col2 = st.columns(2)
            
            with col1:
                translate_on = st.checkbox("翻译文档")
                lang = st.selectbox("目标语言", ["英语", "日语", "德语"], disabled=not translate_on)
                
                analysis_on = st.checkbox("生成专业解读")
                analysis_type = st.selectbox(
                    "解读类型",
                    ["行业研究报告", "公司财报", "市场分析"],
                    disabled=not analysis_on
                )
            
            with col2:
                summary_on = st.checkbox("生成摘要")
                detail_level = st.slider("摘要详细程度", 1, 5, 3, disabled=not summary_on)
                
                if analysis_on:
                    risk_on = st.checkbox("包含风险评估", value=True)
                    suggest_on = st.checkbox("包含投资建议", value=True)

        # 处理按钮
        if st.button("开始分析", type="primary"):
            process_documents(
                uploaded_files,
                translate_on, 
                summary_on,
                analysis_on,
                {
                    "type": analysis_type,
                    "risk": risk_on,
                    "suggestion": suggest_on
                },
                detail_level
            )

def process_documents(files, translate_on, summary_on, analysis_on, analysis_params, detail_level):
    """处理文档主函数"""
    with st.status("正在分析文档...", expanded=True) as status:
        results = {}
        progress_bar = st.progress(0)
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
            futures = []
            for file in files:
                futures.append(executor.submit(
                    process_single_document,
                    file, 
                    translate_on,
                    summary_on,
                    analysis_on,
                    analysis_params,
                    detail_level
                ))
            
            for i, future in enumerate(concurrent.futures.as_completed(futures)):
                try:
                    file_name, outputs = future.result()
                    results[file_name] = outputs
                    progress_bar.progress((i+1)/len(files))
                except Exception as e:
                    st.error(f"处理失败: {str(e)}")
        
        status.update(label="分析完成", state="complete")
    
    # 显示结果
    if results:
        st.subheader("📁 处理结果")
        for file_name, outputs in results.items():
            with st.expander(f"📄 {file_name}", expanded=True):
                for action, path in outputs.items():
                    with open(path, "rb") as f:
                        st.download_button(
                            label=f"下载 {action}",
                            data=f.read(),
                            file_name=os.path.basename(path),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

def process_single_document(file, translate_on, summary_on, analysis_on, analysis_params, detail_level):
    """处理单个文档"""
    outputs = {}
    text = extract_pdf_text(file)
    
    if not text:
        return file.name, outputs
    
    # 翻译处理
    if translate_on:
        prompt = f"将以下内容准确翻译为英文，保留专业术语和数字精度：\n\n{text[:15000]}"
        translated = get_ai_response(prompt)
        if translated:
            trans_path = generate_word_file(translated, file.name, "翻译文档")
            if trans_path:
                outputs["翻译文档"] = trans_path
    
    # 摘要生成
    if summary_on:
        prompt = f'''生成结构化摘要（详细程度{detail_level}/5）：
        - 使用分点列表呈现
        - 关键数据用【】标注
        - 包含核心结论和风险提示
        文档内容：{text[:15000]}'''
        summary = get_ai_response(prompt)
        if summary:
            summary_path = generate_word_file(summary, file.name, "内容摘要")
            if summary_path:
                outputs["内容摘要"] = summary_path
    
    # 研究报告解读
    if analysis_on and text:
        prompt = f'''
        作为{analysis_params['type']}分析师，请从以下维度解读文档：
        1. 核心观点总结（不超过3条）
        2. 关键数据透视（表格呈现）
        { "3. 风险评估（★★★☆）" if analysis_params['risk'] else ""}
        { "4. 投资建议（明确买卖评级）" if analysis_params['suggestion'] else ""}
        5. 后续关注重点
        
        文档名称：《{file.name}》
        内容：{text[:15000]}
        '''
        analysis = get_ai_response(prompt)
        if analysis:
            analysis_path = generate_word_file(analysis, file.name, "专业解读")
            if analysis_path:
                outputs["研究报告解读"] = analysis_path
    
    return file.name, outputs

# ========================
# 主程序
# ========================
def main():
    st.set_page_config(
        page_title="东吴智研金融助手",
        page_icon="📈",
        layout="wide"
    )
    inject_custom_css()
    
    # 安全提示
    st.markdown("""
    <div class="security-alert">
        ⚠️ 安全提示：当前为本地测试模式，请勿在生产环境中使用此配置！
    </div>
    """, unsafe_allow_html=True)
    
    # 页眉
    st.markdown("""
    <div class="header">
        <h1>东吴智研 AI 平台</h1>
        <p>专业金融分析与文档处理系统</p>
    </div>
    """, unsafe_allow_html=True)
    
    # 导航菜单
    page = st.sidebar.radio(
        "功能导航",
        ["💬 智能对话", "📁 文档分析"],
        label_visibility="collapsed"
    )
    
    if page == "💬 智能对话":
        show_chat_interface()
    else:
        show_document_interface()

if __name__ == "__main__":
    main()
